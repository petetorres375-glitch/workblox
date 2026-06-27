import io
import os
import re as _re

from fpdf import FPDF
import docx as python_docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.ats_engine import grade

_FONTS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "fonts")

C_INK    = (27,  34,  48)
C_MID    = (90,  100, 120)
C_SOFT   = (138, 147, 166)
C_ACCENT = (37,  99,  235)
C_GREEN  = (22,  163, 74)
C_RED    = (220, 38,  38)
C_AMBER  = (217, 119, 6)
C_RULE   = (221, 227, 237)
C_BG     = (240, 244, 248)

BG_GREEN = (220, 252, 231)
BG_RED   = (254, 226, 226)
BG_AMBER = (254, 243, 199)
BG_BLUE  = (239, 246, 255)

L_MARGIN  = 15
R_MARGIN  = 15
PAGE_W    = 210
CONTENT_W = PAGE_W - L_MARGIN - R_MARGIN

FONT_NAME = 'DejaVu'
FONT_REG  = os.path.join(_FONTS_DIR, 'DejaVuSans.ttf')
FONT_BOLD = os.path.join(_FONTS_DIR, 'DejaVuSans-Bold.ttf')

_HEADER_RE = _re.compile(
    r'^\s*(?:PROFESSIONAL\s+)?'
    r'(?:SUMMARY|OBJECTIVE|PROFILE|EXPERIENCE|EDUCATION|SKILLS?|CERTIF\w*|ACHIEVE\w*|VOLUNTEER|WORK\s+HISTORY|EMPLOYMENT)\s*$',
    _re.I,
)
_RULE_RE   = _re.compile(r'^[-=]{3,}\s*$')
_BULLET_RE = _re.compile(r'^[•·\-\*]\s+')
_NOTE_RE   = _re.compile(r'^\[\s*REVISION NOTE', _re.I)


def _score_color(score):
    if score >= 65: return C_GREEN
    if score >= 50: return C_AMBER
    return C_RED

def _rgb(triple):
    return RGBColor(*triple)

def _line_type(line):
    s = line.strip()
    if not s:                 return 'blank'
    if _RULE_RE.match(s):    return 'rule'
    if _NOTE_RE.match(s):    return 'note'
    if _HEADER_RE.match(s):  return 'header'
    if s.isupper() and len(s) < 50 and not _re.search(r'[,;.!?]', s):
        return 'header'
    if _BULLET_RE.match(s):  return 'bullet'
    return 'body'


class _PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font(FONT_NAME, '',  FONT_REG)
        self.add_font(FONT_NAME, 'B', FONT_BOLD)

    def section_title(self, title):
        self.ln(2)
        self.set_font(FONT_NAME, 'B', 7.5)
        self.set_text_color(*C_SOFT)
        self.cell(CONTENT_W, 5, title.upper(), ln=False)
        self.ln(5)
        self.set_draw_color(*C_RULE)
        self.set_line_width(0.3)
        self.line(L_MARGIN, self.get_y(), L_MARGIN + CONTENT_W, self.get_y())
        self.ln(4)
        self.set_text_color(*C_INK)
        self.set_draw_color(0, 0, 0)

    def chips(self, keywords, found):
        if not keywords:
            return
        self.set_font(FONT_NAME, '', 8)
        bg = BG_GREEN if found else BG_RED
        fg = C_GREEN  if found else C_RED
        for kw in keywords:
            text = f" {kw} "
            w = self.get_string_width(text) + 2
            h = 5.5
            x, y = self.get_x(), self.get_y()
            if x + w > L_MARGIN + CONTENT_W:
                self.set_xy(L_MARGIN, y + h + 1)
                x, y = self.get_x(), self.get_y()
            self.set_fill_color(*bg)
            self.set_draw_color(*bg)
            self.rect(x, y, w, h, 'F')
            self.set_text_color(*fg)
            self.set_xy(x, y)
            self.cell(w, h, text, ln=False)
            self.set_xy(x + w + 2, y)
        self.set_text_color(*C_INK)
        self.set_draw_color(0, 0, 0)
        self.ln(8)

    def bar(self, pct, color=C_ACCENT, height=4):
        y = self.get_y()
        self.set_fill_color(*C_RULE)
        self.rect(L_MARGIN, y, CONTENT_W, height, 'F')
        if pct > 0:
            self.set_fill_color(*color)
            self.rect(L_MARGIN, y, CONTENT_W * (pct / 100), height, 'F')
        self.ln(height + 3)


def generate_pdf(results, client_name, filename, job_role, now):
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(L_MARGIN, 15, R_MARGIN)

    score = results['score']
    sc    = _score_color(score)

    pdf.set_fill_color(*C_INK)
    pdf.rect(0, 0, PAGE_W, 20, 'F')
    pdf.set_xy(L_MARGIN, 6)
    pdf.set_font(FONT_NAME, 'B', 10)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(CONTENT_W, 8, 'ATS RESUME ANALYZER  —  REPORT')
    pdf.ln(24)

    pdf.set_font(FONT_NAME, 'B', 18)
    pdf.set_text_color(*C_INK)
    pdf.cell(CONTENT_W, 9, client_name, ln=True)
    pdf.set_font(FONT_NAME, '', 8.5)
    pdf.set_text_color(*C_SOFT)
    meta = filename + (f'   |   Role: {job_role}' if job_role else '') + f'   |   {now}'
    pdf.cell(CONTENT_W, 5, meta, ln=True)
    pdf.ln(5)

    block_y = pdf.get_y()
    pdf.set_fill_color(*C_BG)
    pdf.rect(L_MARGIN, block_y, CONTENT_W, 30, 'F')

    pdf.set_xy(L_MARGIN + 4, block_y + 3)
    pdf.set_font(FONT_NAME, 'B', 34)
    pdf.set_text_color(*sc)
    pdf.cell(26, 14, str(score), ln=False)

    pdf.set_xy(L_MARGIN + 30, block_y + 9)
    pdf.set_font(FONT_NAME, '', 8.5)
    pdf.set_text_color(*C_SOFT)
    pdf.cell(15, 5, '/ 100', ln=False)

    pdf.set_xy(L_MARGIN + 54, block_y + 3)
    pdf.set_font(FONT_NAME, 'B', 11)
    pdf.set_text_color(*C_INK)
    pdf.cell(100, 7, grade(score), ln=False)

    pdf.set_xy(L_MARGIN + 54, block_y + 11)
    pdf.set_font(FONT_NAME, '', 8.5)
    pdf.set_text_color(*C_MID)
    pdf.cell(100, 5, f"{results['total_found']} / {results['total_possible']} keywords found", ln=False)

    bar_x, bar_y, bar_w = L_MARGIN + 54, block_y + 20, CONTENT_W - 58
    pdf.set_fill_color(*C_RULE)
    pdf.rect(bar_x, bar_y, bar_w, 4, 'F')
    pdf.set_fill_color(*sc)
    pdf.rect(bar_x, bar_y, bar_w * (score / 100), 4, 'F')
    pdf.set_xy(L_MARGIN, block_y + 32)
    pdf.ln(4)

    pdf.section_title('Contact Information')
    contact = results.get('contact', {})
    for label, key in [('Email', 'email'), ('Phone', 'phone'), ('LinkedIn', 'linkedin'), ('Location', 'location')]:
        present = contact.get(key, False)
        bg, fg  = (BG_GREEN, C_GREEN) if present else (BG_RED, C_RED)
        text = f"  {'OK' if present else 'MISSING'}  {label}  "
        w = pdf.get_string_width(text) + 2
        x, y = pdf.get_x(), pdf.get_y()
        pdf.set_fill_color(*bg); pdf.set_draw_color(*bg)
        pdf.rect(x, y, w, 6, 'F')
        pdf.set_font(FONT_NAME, '', 8.5); pdf.set_text_color(*fg)
        pdf.set_xy(x, y); pdf.cell(w, 6, text, ln=False)
        pdf.set_xy(x + w + 3, y)
    pdf.set_text_color(*C_INK); pdf.set_draw_color(0, 0, 0)
    pdf.ln(10)

    recs = results.get('recommendations', [])
    if recs:
        pdf.section_title('Recommendations')
        p_colors = {'high': (C_RED, BG_RED), 'medium': (C_AMBER, BG_AMBER), 'low': (C_ACCENT, BG_BLUE)}
        for rec in recs:
            fg, bg = p_colors.get(rec['priority'], (C_MID, C_BG))
            x, y   = pdf.get_x(), pdf.get_y()
            badge  = f"  {rec['priority'].upper()}  "
            pdf.set_font(FONT_NAME, 'B', 7)
            bw = pdf.get_string_width(badge) + 2
            pdf.set_fill_color(*bg); pdf.set_draw_color(*bg)
            pdf.rect(x, y + 0.5, bw, 5.5, 'F')
            pdf.set_text_color(*fg); pdf.set_xy(x, y)
            pdf.cell(bw, 6.5, badge, ln=False)
            pdf.set_font(FONT_NAME, 'B', 9); pdf.set_text_color(*C_INK)
            pdf.cell(CONTENT_W - bw, 6.5, f"  {rec['title']}", ln=True)
            pdf.set_font(FONT_NAME, '', 8.5); pdf.set_text_color(*C_MID)
            pdf.set_x(L_MARGIN + 4)
            pdf.multi_cell(CONTENT_W - 4, 5, rec['detail'], ln=True)
            pdf.set_draw_color(0, 0, 0); pdf.ln(1)
        pdf.ln(2)

    if results.get('warnings'):
        pdf.section_title('Warnings')
        for msg in results['warnings']:
            pdf.set_fill_color(*BG_AMBER); pdf.set_draw_color(*BG_AMBER)
            y = pdf.get_y()
            pdf.rect(L_MARGIN, y, CONTENT_W, 7, 'F')
            pdf.set_font(FONT_NAME, '', 9); pdf.set_text_color(146, 64, 14)
            pdf.set_xy(L_MARGIN, y); pdf.cell(CONTENT_W, 7, f"  ! {msg}", ln=True)
            pdf.ln(1)
        pdf.set_text_color(*C_INK); pdf.set_draw_color(0, 0, 0); pdf.ln(2)

    if results.get('tips'):
        pdf.section_title('Tips')
        for tip in results['tips']:
            pdf.set_font(FONT_NAME, '', 9); pdf.set_text_color(*C_MID)
            pdf.cell(CONTENT_W, 5.5, f"  -> {tip}", ln=True)
        pdf.set_text_color(*C_INK); pdf.ln(3)

    q = results.get('quantification', {})
    pdf.section_title('Measurable Results')
    if q.get('has_metrics'):
        pdf.set_font(FONT_NAME, 'B', 9); pdf.set_text_color(*C_GREEN)
        pdf.cell(CONTENT_W, 5, f"Metrics detected  ({q.get('number_count', 0)} numbers found)", ln=True)
        if q.get('examples'):
            pdf.set_font(FONT_NAME, '', 8.5); pdf.set_text_color(*C_MID)
            pdf.cell(CONTENT_W, 5, 'Examples: ' + '  |  '.join(str(e) for e in q['examples']), ln=True)
    else:
        pdf.set_font(FONT_NAME, 'B', 9); pdf.set_text_color(*C_RED)
        pdf.cell(CONTENT_W, 5, 'No metrics found', ln=True)
        pdf.set_font(FONT_NAME, '', 8.5); pdf.set_text_color(*C_MID)
        pdf.multi_cell(CONTENT_W, 5, 'Add numbers to bullets — e.g. "Managed 12-person team"', ln=True)
    pdf.set_text_color(*C_INK); pdf.ln(3)

    pdf.section_title('Resume Sections')
    sections = results.get('sections', {})
    col, col_w = 0, CONTENT_W / 2
    for name, present in sections.items():
        pdf.set_font(FONT_NAME, '', 9)
        pdf.set_text_color(*C_GREEN if present else C_SOFT)
        mark = '[+]' if present else '[ ]'
        if col == 0: pdf.set_x(L_MARGIN)
        pdf.cell(col_w, 6, f"  {mark}  {name}", ln=(col == 1))
        col = 1 - col
    if col == 1: pdf.ln(6)
    pdf.set_text_color(*C_INK); pdf.ln(3)

    if results.get('job_match'):
        jm = results['job_match']
        pdf.section_title(f"Job Match: {jm['role'].upper()}")
        pdf.set_font(FONT_NAME, 'B', 12); pdf.set_text_color(*C_ACCENT)
        pdf.cell(CONTENT_W, 7, f"{jm['score']} / {jm['total']} keywords matched", ln=True)
        pdf.set_text_color(*C_INK); pdf.ln(2)
        if jm['found']:
            pdf.set_font(FONT_NAME, 'B', 7.5); pdf.set_text_color(*C_SOFT)
            pdf.cell(CONTENT_W, 4.5, 'FOUND', ln=True)
            pdf.chips(jm['found'], True)
        if jm['missing']:
            pdf.set_font(FONT_NAME, 'B', 7.5); pdf.set_text_color(*C_SOFT)
            pdf.cell(CONTENT_W, 4.5, 'MISSING — add these to the resume', ln=True)
            pdf.chips(jm['missing'], False)

    pdf.section_title('Keyword Breakdown by Category')
    pdf.ln(1)
    for cat, data in results['categories'].items():
        pct = int((data['score'] / data['total']) * 100) if data['total'] else 0
        pdf.set_font(FONT_NAME, 'B', 9.5); pdf.set_text_color(*C_INK)
        pdf.cell(CONTENT_W, 5.5, f"{cat}  —  {data['score']} / {data['total']}  ({pct}%)", ln=True)
        pdf.bar(pct)
        if data['found']:
            pdf.set_font(FONT_NAME, 'B', 7.5); pdf.set_text_color(*C_SOFT)
            pdf.cell(CONTENT_W, 4.5, 'FOUND', ln=True)
            pdf.chips(data['found'], True)
        if data['missing']:
            pdf.set_font(FONT_NAME, 'B', 7.5); pdf.set_text_color(*C_SOFT)
            pdf.cell(CONTENT_W, 4.5, 'MISSING', ln=True)
            pdf.chips(data['missing'], False)
        pdf.ln(2)

    pdf.set_y(-14)
    pdf.set_font(FONT_NAME, '', 7.5); pdf.set_text_color(*C_SOFT)
    pdf.cell(CONTENT_W, 5, f"Generated by Workblox — Torres Tech Remote  |  {now}", align='C')

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def generate_docx(results, client_name, filename, job_role, now):
    doc = python_docx.Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.75)
        section.left_margin = section.right_margin = Inches(0.9)

    score = results['score']
    sc    = _rgb(_score_color(score))

    def _heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(title.upper())
        r.font.name = 'Calibri'; r.font.size = Pt(7.5)
        r.font.bold = True; r.font.color.rgb = _rgb(C_SOFT)

    def _kw_para(keywords, found):
        if not keywords: return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        fg = _rgb(C_GREEN) if found else _rgb(C_RED)
        prefix = '+ ' if found else '- '
        for i, kw in enumerate(keywords):
            if i > 0:
                r = p.add_run('  '); r.font.size = Pt(9)
            r = p.add_run(prefix + kw)
            r.font.size = Pt(9); r.font.color.rgb = fg

    r = doc.add_paragraph().add_run('ATS Resume Analyzer — Report')
    r.font.name = 'Calibri'; r.font.size = Pt(9); r.font.color.rgb = _rgb(C_SOFT)

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(client_name)
    r.font.name = 'Calibri'; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = _rgb(C_INK)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(8)
    meta_text = filename + (f'   |   Role: {job_role}' if job_role else '') + f'   |   {now}'
    r = meta_p.add_run(meta_text)
    r.font.size = Pt(8.5); r.font.color.rgb = _rgb(C_SOFT)

    _heading('Overall Score')
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    r = sp.add_run(f'{score} / 100')
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = sc
    r2 = sp.add_run(f'   {grade(score)}')
    r2.font.size = Pt(13); r2.font.bold = True; r2.font.color.rgb = _rgb(C_INK)

    _heading('Contact Information')
    contact = results.get('contact', {})
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'; table.autofit = True
    for i, (label, key) in enumerate([('Email', 'email'), ('Phone', 'phone'), ('LinkedIn', 'linkedin'), ('Location', 'location')]):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        present = contact.get(key, False)
        r = cell.paragraphs[0].add_run(f"{'YES' if present else 'MISSING'}  {label}")
        r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = _rgb(C_GREEN) if present else _rgb(C_RED)

    recs = results.get('recommendations', [])
    if recs:
        _heading('Recommendations')
        p_map = {'high': C_RED, 'medium': C_AMBER, 'low': C_ACCENT}
        for rec in recs:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
            br = p.add_run(f"[{rec['priority'].upper()}]  ")
            br.font.size = Pt(8); br.font.bold = True; br.font.color.rgb = _rgb(p_map.get(rec['priority'], C_MID))
            tr = p.add_run(rec['title'])
            tr.font.size = Pt(10); tr.font.bold = True; tr.font.color.rgb = _rgb(C_INK)
            dp = doc.add_paragraph(); dp.paragraph_format.space_after = Pt(5)
            dp.paragraph_format.left_indent = Inches(0.25)
            dr = dp.add_run(rec['detail'])
            dr.font.size = Pt(9); dr.font.color.rgb = _rgb(C_MID)

    if results.get('warnings'):
        _heading('Warnings')
        for msg in results['warnings']:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"!  {msg}")
            r.font.size = Pt(9.5); r.font.color.rgb = _rgb(C_AMBER)

    _heading('Keyword Breakdown by Category')
    for cat, data in results['categories'].items():
        pct = int((data['score'] / data['total']) * 100) if data['total'] else 0
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{cat}  —  {data['score']} / {data['total']}  ({pct}%)")
        r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = _rgb(C_INK)
        if data['found']:   _kw_para(data['found'], True)
        if data['missing']: _kw_para(data['missing'], False)

    if results.get('job_match'):
        jm = results['job_match']
        _heading(f"Job Match: {jm['role']}")
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{jm['score']} / {jm['total']} keywords matched")
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = _rgb(C_ACCENT)
        if jm['found']:   _kw_para(jm['found'], True)
        if jm['missing']: _kw_para(jm['missing'], False)

    doc.add_paragraph()
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f"Generated by Workblox — Torres Tech Remote  |  {now}")
    r.font.size = Pt(8); r.font.color.rgb = _rgb(C_SOFT)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
