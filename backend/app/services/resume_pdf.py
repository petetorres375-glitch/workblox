import io
import os

from fpdf import FPDF
import docx as python_docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_FONTS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "fonts")

C_NAME    = (15,  23,  42)
C_ACCENT  = (37,  99,  235)
C_BODY    = (30,  30,  30)
C_MUTED   = (90, 100, 120)
C_RULE    = (200, 210, 225)

FONT_NAME = "DejaVu"
FONT_REG  = os.path.join(_FONTS_DIR, "DejaVuSans.ttf")
FONT_BOLD = os.path.join(_FONTS_DIR, "DejaVuSans-Bold.ttf")

L_MARGIN  = 18
R_MARGIN  = 18
PAGE_W    = 210
CONTENT_W = PAGE_W - L_MARGIN - R_MARGIN


class _PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font(FONT_NAME, "",  FONT_REG)
        self.add_font(FONT_NAME, "B", FONT_BOLD)

    def section_header(self, title):
        self.ln(4)
        self.set_font(FONT_NAME, "B", 9)
        self.set_text_color(*C_ACCENT)
        self.cell(CONTENT_W, 5, title.upper(), ln=True)
        self.set_draw_color(*C_RULE)
        self.set_line_width(0.4)
        self.line(L_MARGIN, self.get_y(), L_MARGIN + CONTENT_W, self.get_y())
        self.ln(3)
        self.set_text_color(*C_BODY)
        self.set_draw_color(0, 0, 0)

    def bullet(self, text):
        self.set_font(FONT_NAME, "", 9.5)
        self.set_x(L_MARGIN + 4)
        self.cell(4, 5.5, "•", ln=False)
        self.set_x(L_MARGIN + 8)
        self.multi_cell(CONTENT_W - 8, 5.5, text.strip(), ln=True)


def generate_resume_pdf(contact, resume, job_role):
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(L_MARGIN, 16, R_MARGIN)

    # Name
    pdf.set_font(FONT_NAME, "B", 22)
    pdf.set_text_color(*C_NAME)
    pdf.cell(CONTENT_W, 10, contact.get("name", ""), ln=True)

    # Contact line
    parts = []
    for key in ("email", "phone", "location", "linkedin"):
        val = contact.get(key, "").strip()
        if val:
            parts.append(val)
    pdf.set_font(FONT_NAME, "", 8.5)
    pdf.set_text_color(*C_MUTED)
    pdf.cell(CONTENT_W, 5, "  |  ".join(parts), ln=True)

    if job_role:
        pdf.set_font(FONT_NAME, "B", 8.5)
        pdf.set_text_color(*C_ACCENT)
        pdf.cell(CONTENT_W, 5, job_role.title(), ln=True)

    pdf.ln(1)
    pdf.set_draw_color(*C_NAME)
    pdf.set_line_width(0.6)
    pdf.line(L_MARGIN, pdf.get_y(), L_MARGIN + CONTENT_W, pdf.get_y())
    pdf.ln(4)

    # Summary
    if resume.get("summary"):
        pdf.section_header("Professional Summary")
        pdf.set_font(FONT_NAME, "", 9.5)
        pdf.set_text_color(*C_BODY)
        pdf.multi_cell(CONTENT_W, 5.5, resume["summary"], ln=True)
        pdf.ln(1)

    # Experience
    exp = resume.get("experience", [])
    if exp:
        pdf.section_header("Work Experience")
        for job in exp:
            pdf.set_font(FONT_NAME, "B", 10)
            pdf.set_text_color(*C_NAME)
            title_w = pdf.get_string_width(job.get("title", "")) + 2
            pdf.cell(title_w, 6, job.get("title", ""), ln=False)
            pdf.set_font(FONT_NAME, "", 9)
            pdf.set_text_color(*C_MUTED)
            pdf.cell(4, 6, "  —", ln=False)
            pdf.cell(0, 6, f"  {job.get('company', '')}", ln=False)
            dates = job.get("dates", "")
            if dates:
                pdf.set_x(L_MARGIN + CONTENT_W - pdf.get_string_width(dates))
                pdf.cell(pdf.get_string_width(dates), 6, dates, ln=True)
            else:
                pdf.ln(6)
            for bullet in job.get("bullets", []):
                pdf.bullet(bullet)
            pdf.ln(2)

    # Education
    edu = resume.get("education", [])
    if edu:
        pdf.section_header("Education")
        for e in edu:
            pdf.set_font(FONT_NAME, "B", 10)
            pdf.set_text_color(*C_NAME)
            pdf.cell(CONTENT_W * 0.75, 6, e.get("degree", ""), ln=False)
            pdf.set_font(FONT_NAME, "", 9)
            pdf.set_text_color(*C_MUTED)
            year = e.get("year", "")
            if year:
                pdf.set_x(L_MARGIN + CONTENT_W - pdf.get_string_width(year))
                pdf.cell(pdf.get_string_width(year), 6, year, ln=True)
            else:
                pdf.ln(6)
            pdf.set_font(FONT_NAME, "", 9)
            pdf.set_text_color(*C_MUTED)
            pdf.cell(CONTENT_W, 5, e.get("school", ""), ln=True)
            pdf.ln(1)

    # Skills
    skills = resume.get("skills", [])
    if skills:
        pdf.section_header("Skills")
        pdf.set_font(FONT_NAME, "", 9.5)
        pdf.set_text_color(*C_BODY)
        pdf.multi_cell(CONTENT_W, 5.5, " • ".join(skills), ln=True)
        pdf.ln(1)

    # Certifications
    certs = resume.get("certifications", [])
    if certs:
        pdf.section_header("Certifications")
        for cert in certs:
            if cert.strip():
                pdf.bullet(cert)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def generate_resume_docx(contact, resume, job_role):
    doc = python_docx.Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    def _rgb(t): return RGBColor(*t)

    def _section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(title.upper())
        r.font.name = "Calibri"; r.font.size = Pt(9)
        r.font.bold = True; r.font.color.rgb = _rgb(C_ACCENT)

    # Name
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(1)
    r = name_p.add_run(contact.get("name", ""))
    r.font.name = "Calibri"; r.font.size = Pt(22)
    r.font.bold = True; r.font.color.rgb = _rgb(C_NAME)

    # Contact
    parts = [contact.get(k, "").strip() for k in ("email", "phone", "location", "linkedin") if contact.get(k, "").strip()]
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(1)
    r = contact_p.add_run("  |  ".join(parts))
    r.font.size = Pt(8.5); r.font.color.rgb = _rgb(C_MUTED)

    if job_role:
        role_p = doc.add_paragraph()
        role_p.paragraph_format.space_after = Pt(4)
        r = role_p.add_run(job_role.title())
        r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = _rgb(C_ACCENT)

    if resume.get("summary"):
        _section_heading("Professional Summary")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(resume["summary"])
        r.font.size = Pt(10); r.font.color.rgb = _rgb(C_BODY)

    exp = resume.get("experience", [])
    if exp:
        _section_heading("Work Experience")
        for job in exp:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(job.get("title", ""))
            r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = _rgb(C_NAME)
            r2 = p.add_run(f"   {job.get('company', '')}   {job.get('dates', '')}")
            r2.font.size = Pt(9); r2.font.color.rgb = _rgb(C_MUTED)
            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.25)
                r = bp.add_run(bullet.strip())
                r.font.size = Pt(10); r.font.color.rgb = _rgb(C_BODY)

    edu = resume.get("education", [])
    if edu:
        _section_heading("Education")
        for e in edu:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(e.get("degree", ""))
            r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = _rgb(C_NAME)
            r2 = p.add_run(f"   {e.get('school', '')}   {e.get('year', '')}")
            r2.font.size = Pt(9); r2.font.color.rgb = _rgb(C_MUTED)

    skills = resume.get("skills", [])
    if skills:
        _section_heading("Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(" • ".join(skills))
        r.font.size = Pt(10); r.font.color.rgb = _rgb(C_BODY)

    certs = resume.get("certifications", [])
    if certs:
        _section_heading("Certifications")
        for cert in certs:
            if cert.strip():
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(cert.strip())
                r.font.size = Pt(10); r.font.color.rgb = _rgb(C_BODY)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_resume_txt(contact, resume, job_role):
    lines = []
    name = contact.get("name", "")
    if name:
        lines += [name, "=" * len(name)]

    parts = [contact.get(k, "").strip() for k in ("email", "phone", "location", "linkedin") if contact.get(k, "").strip()]
    if parts:
        lines.append("  |  ".join(parts))
    if job_role:
        lines.append(job_role.title())
    lines.append("")

    if resume.get("summary"):
        lines += ["PROFESSIONAL SUMMARY", "-" * 30, resume["summary"], ""]

    if resume.get("experience"):
        lines += ["WORK EXPERIENCE", "-" * 30]
        for job in resume["experience"]:
            lines.append(f"{job.get('title', '')}  —  {job.get('company', '')}  |  {job.get('dates', '')}")
            for b in job.get("bullets", []):
                lines.append(f"  • {b.strip()}")
            lines.append("")

    if resume.get("education"):
        lines += ["EDUCATION", "-" * 30]
        for e in resume["education"]:
            lines.append(f"{e.get('degree', '')}  |  {e.get('school', '')}  |  {e.get('year', '')}")
        lines.append("")

    if resume.get("skills"):
        lines += ["SKILLS", "-" * 30, " • ".join(resume["skills"]), ""]

    if resume.get("certifications"):
        lines += ["CERTIFICATIONS", "-" * 30]
        for c in resume["certifications"]:
            if c.strip():
                lines.append(f"  • {c.strip()}")

    return "\n".join(lines)
